"""Document processing service - PDF, DOCX parsing and chunking."""
import io
from sqlalchemy.orm import Session
from sqlalchemy import select

from backend.models.knowledge import Document, DocumentChunk
from backend.services import embeddings
from backend.core.logger import log_upload

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def _extract_text_pdf(content: bytes) -> str:
    """Extract text from PDF."""
    import fitz  # PyMuPDF
    
    doc = fitz.open(stream=content, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def _extract_text_docx(content: bytes) -> str:
    """Extract text from DOCX."""
    from docx import Document as DocxDocument
    
    doc = DocxDocument(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks."""
    text = text.strip()
    if len(text) <= chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at sentence/paragraph boundary
        if end < len(text):
            for sep in ["\n\n", "\n", ". ", "! ", "? "]:
                last_sep = chunk.rfind(sep)
                if last_sep > chunk_size // 2:
                    chunk = chunk[:last_sep + len(sep)]
                    end = start + len(chunk)
                    break
        
        chunks.append(chunk.strip())
        start = end - overlap
    
    return [c for c in chunks if c]


def upload(db: Session, agent_id: int, filename: str, content: bytes) -> Document:
    """Upload and process a document."""
    ext = filename.rsplit(".", 1)[-1].lower()
    
    if ext == "pdf":
        text = _extract_text_pdf(content)
    elif ext in ("docx", "doc"):
        text = _extract_text_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
    chunks = _chunk_text(text)
    if not chunks:
        raise ValueError("No text content found in document")
    
    chunk_embeddings = embeddings.get_embeddings_batch(chunks)
    
    doc = Document(
        agent_id=agent_id,
        filename=filename,
        file_type=ext,
        file_size=len(content),
        chunk_count=len(chunks)
    )
    db.add(doc)
    db.flush()
    
    for i, (chunk_text, chunk_emb) in enumerate(zip(chunks, chunk_embeddings)):
        chunk = DocumentChunk(
            document_id=doc.id,
            content=chunk_text,
            chunk_index=i,
            embedding=chunk_emb
        )
        db.add(chunk)
    
    db.commit()
    db.refresh(doc)
    
    log_upload("document", filename, f"{len(chunks)} chunks")
    return doc


def delete(db: Session, doc_id: int) -> bool:
    doc = db.get(Document, doc_id)
    if not doc:
        return False
    db.delete(doc)
    db.commit()
    return True


def get_by_agent(db: Session, agent_id: int) -> list[Document]:
    """Get all documents for an agent."""
    return list(db.scalars(
        select(Document)
        .where(Document.agent_id == agent_id, Document.is_active == True)
        .order_by(Document.created_at.desc())
    ))


def search(db: Session, agent_id: int, query: str, limit: int = 5) -> list[dict]:
    """Semantic search across document chunks."""
    query_embedding = embeddings.get_embedding(query)
    
    results = db.execute(
        select(DocumentChunk, Document)
        .join(Document)
        .where(
            Document.agent_id == agent_id,
            Document.is_active == True
        )
        .order_by(DocumentChunk.embedding.cosine_distance(query_embedding))
        .limit(limit)
    ).all()
    
    return [
        {
            "document": row.Document.filename,
            "content": row.DocumentChunk.content,
            "chunk_index": row.DocumentChunk.chunk_index
        }
        for row in results
    ]
